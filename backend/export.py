import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel

from summarization import ActionItem

_RTL_LANGUAGES: frozenset[str] = frozenset({
    'he', 'heb', 'hebrew', 'עברית',
    'ar', 'ara', 'arabic', 'عربي', 'عربية',
    'fa', 'per', 'fas', 'persian', 'farsi', 'فارسی',
    'ur', 'urd', 'urdu', 'اردو',
    'yi', 'yid', 'yiddish',
    'pashto', 'sindhi',
})


class ExportRequest(BaseModel):
    language: str
    summary: str
    participants: list[str]
    decisions: list[str]
    action_items: list[ActionItem]
    transcript: str


def _set_paragraph_rtl(paragraph) -> None:
    """Insert w:bidi into paragraph properties to enable RTL text direction."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.insert(0, bidi)


def _set_table_rtl(table) -> None:
    """Set w:bidiVisual on the table so Word renders column order right-to-left."""
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    bidiVisual = OxmlElement('w:bidiVisual')
    tblPr.append(bidiVisual)


def _apply_rtl_to_cells(cells) -> None:
    """Apply RTL to every paragraph in each cell of a table row."""
    for cell in cells:
        for paragraph in cell.paragraphs:
            _set_paragraph_rtl(paragraph)


def build_docx(data: ExportRequest) -> bytes:
    is_rtl = data.language.lower().strip() in _RTL_LANGUAGES
    doc = Document()

    doc.add_heading('Meeting Summary', level=0)

    doc.add_heading('Overview', level=1)
    p = doc.add_paragraph(data.summary)
    if is_rtl:
        _set_paragraph_rtl(p)

    doc.add_heading('Participants', level=1)
    if data.participants:
        for name in data.participants:
            p = doc.add_paragraph(name, style='List Bullet')
            if is_rtl:
                _set_paragraph_rtl(p)
    else:
        p = doc.add_paragraph('No participants identified.')
        if is_rtl:
            _set_paragraph_rtl(p)

    doc.add_heading('Decisions', level=1)
    if data.decisions:
        for decision in data.decisions:
            p = doc.add_paragraph(decision, style='List Bullet')
            if is_rtl:
                _set_paragraph_rtl(p)
    else:
        p = doc.add_paragraph('No decisions recorded.')
        if is_rtl:
            _set_paragraph_rtl(p)

    doc.add_heading('Action Items', level=1)
    if data.action_items:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        if is_rtl:
            _set_table_rtl(table)
        hdr = table.rows[0].cells
        hdr[0].text = 'Task'
        hdr[1].text = 'Owner'
        hdr[2].text = 'Due'
        if is_rtl:
            _apply_rtl_to_cells(hdr)
        for item in data.action_items:
            row = table.add_row().cells
            row[0].text = item.task
            row[1].text = item.owner
            row[2].text = item.due or '—'
            if is_rtl:
                _apply_rtl_to_cells(row)
    else:
        p = doc.add_paragraph('No action items.')
        if is_rtl:
            _set_paragraph_rtl(p)

    doc.add_heading('Full Transcript', level=1)
    p = doc.add_paragraph(data.transcript)
    if is_rtl:
        _set_paragraph_rtl(p)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
