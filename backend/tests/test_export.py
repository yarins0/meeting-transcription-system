import io

import pytest
from docx import Document

from export import ExportRequest, build_docx
from summarization import ActionItem


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _ltr_request(**overrides) -> ExportRequest:
    base = dict(
        language="English",
        summary="A productive meeting.",
        participants=["Alice", "Bob"],
        decisions=["Proceed with plan A."],
        action_items=[ActionItem(task="Review the draft", owner="Alice", due="2024-01-15")],
        transcript="Alice: Hello.\nBob: Hi there.",
    )
    base.update(overrides)
    return ExportRequest(**base)


def _parse_docx(docx_bytes: bytes) -> Document:
    return Document(io.BytesIO(docx_bytes))


def _all_paragraph_texts(doc: Document) -> list[str]:
    return [p.text for p in doc.paragraphs]


# ---------------------------------------------------------------------------
# Section presence
# ---------------------------------------------------------------------------

def test_build_docx_returns_bytes():
    result = build_docx(_ltr_request())
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_ltr_document_contains_all_sections():
    doc = _parse_docx(build_docx(_ltr_request()))
    texts = _all_paragraph_texts(doc)

    assert "Meeting Summary" in texts
    assert "Overview" in texts
    assert "Participants" in texts
    assert "Decisions" in texts
    assert "Action Items" in texts
    assert "Full Transcript" in texts


def test_summary_text_appears_in_document():
    doc = _parse_docx(build_docx(_ltr_request()))
    texts = _all_paragraph_texts(doc)
    assert "A productive meeting." in texts


def test_participants_appear_in_document():
    doc = _parse_docx(build_docx(_ltr_request()))
    texts = _all_paragraph_texts(doc)
    assert "Alice" in texts
    assert "Bob" in texts


def test_decisions_appear_in_document():
    doc = _parse_docx(build_docx(_ltr_request()))
    texts = _all_paragraph_texts(doc)
    assert "Proceed with plan A." in texts


def test_transcript_appears_in_document():
    doc = _parse_docx(build_docx(_ltr_request()))
    texts = _all_paragraph_texts(doc)
    assert "Alice: Hello.\nBob: Hi there." in texts


# ---------------------------------------------------------------------------
# Action items table
# ---------------------------------------------------------------------------

def test_action_items_table_has_correct_columns():
    doc = _parse_docx(build_docx(_ltr_request()))
    assert len(doc.tables) == 1
    header_row = doc.tables[0].rows[0]
    cell_texts = [c.text for c in header_row.cells]
    assert cell_texts == ["Task", "Owner", "Due"]


def test_action_item_data_row_is_correct():
    doc = _parse_docx(build_docx(_ltr_request()))
    data_row = doc.tables[0].rows[1]
    assert data_row.cells[0].text == "Review the draft"
    assert data_row.cells[1].text == "Alice"
    assert data_row.cells[2].text == "2024-01-15"


def test_null_due_date_renders_as_dash():
    req = _ltr_request(
        action_items=[ActionItem(task="Check logs", owner="Bob", due=None)]
    )
    doc = _parse_docx(build_docx(req))
    data_row = doc.tables[0].rows[1]
    assert data_row.cells[2].text == "—"


# ---------------------------------------------------------------------------
# Empty collections
# ---------------------------------------------------------------------------

def test_empty_decisions_shows_fallback_text():
    doc = _parse_docx(build_docx(_ltr_request(decisions=[])))
    texts = _all_paragraph_texts(doc)
    assert "No decisions recorded." in texts


def test_empty_action_items_shows_fallback_text():
    doc = _parse_docx(build_docx(_ltr_request(action_items=[])))
    texts = _all_paragraph_texts(doc)
    assert "No action items." in texts
    assert len(doc.tables) == 0


def test_empty_participants_shows_fallback_text():
    doc = _parse_docx(build_docx(_ltr_request(participants=[])))
    texts = _all_paragraph_texts(doc)
    assert "No participants identified." in texts


# ---------------------------------------------------------------------------
# RTL support
# ---------------------------------------------------------------------------

def test_hebrew_language_sets_rtl_bidi_on_paragraphs():
    req = _ltr_request(language="hebrew")
    doc = _parse_docx(build_docx(req))
    # The summary paragraph should have w:bidi in its XML.
    summary_para = next(p for p in doc.paragraphs if p.text == "A productive meeting.")
    assert "w:bidi" in summary_para._p.xml


def test_ltr_document_has_no_bidi_on_summary_paragraph():
    doc = _parse_docx(build_docx(_ltr_request(language="English")))
    summary_para = next(p for p in doc.paragraphs if p.text == "A productive meeting.")
    assert "w:bidi" not in summary_para._p.xml


def test_hebrew_table_has_bidi_visual():
    req = _ltr_request(language="hebrew")
    doc = _parse_docx(build_docx(req))
    tbl_xml = doc.tables[0]._tbl.xml
    assert "w:bidiVisual" in tbl_xml


@pytest.mark.parametrize("lang", ["he", "heb", "עברית", "ar", "arabic", "عربية", "fa", "persian"])
def test_rtl_language_variants_all_trigger_rtl(lang: str):
    req = _ltr_request(language=lang)
    doc = _parse_docx(build_docx(req))
    summary_para = next(p for p in doc.paragraphs if p.text == "A productive meeting.")
    assert "w:bidi" in summary_para._p.xml
